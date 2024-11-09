import fitz  # PyMuPDF برای خواندن PDF
from telegram import Update, ReplyKeyboardMarkup
from telegram.ext import Updater, CommandHandler, MessageHandler, Filters, CallbackContext, ConversationHandler
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# دریافت توکن ربات از کاربر به زبان انگلیسی
TOKEN = input("Please enter your Telegram bot token: ")

# مراحل مکالمه
TEXT, FORMAT, FONT_SIZE, FONT_STYLE, FONT_BOLD_ITALIC, EDIT = range(6)

def start(update: Update, context: CallbackContext) -> int:
    update.message.reply_text(
        "Hello! Please send the text you want to convert to a file.\n"
        "To cancel the conversation, enter /cancel."
    )
    return TEXT

def get_text(update: Update, context: CallbackContext) -> int:
    context.user_data['text'] = update.message.text.split('\n')
    reply_keyboard = [['txt', 'docx', 'pdf']]
    update.message.reply_text(
        "Text received. Please choose the output file format (txt, docx, or pdf):",
        reply_markup=ReplyKeyboardMarkup(reply_keyboard, one_time_keyboard=True)
    )
    return FORMAT

def get_format(update: Update, context: CallbackContext) -> int:
    context.user_data['format'] = update.message.text.lower()
    if context.user_data['format'] == 'docx':
        update.message.reply_text(
            "docx format selected. Please choose a font size between 10 and 20 (e.g., 12):"
        )
        return FONT_SIZE
    elif context.user_data['format'] == 'pdf':
        update.message.reply_text("PDF format selected. The text will be converted to a PDF file.")
        convert_to_pdf(update, context)
        return ConversationHandler.END
    elif context.user_data['format'] == 'txt':
        update.message.reply_text("txt format selected. A simple text file will be sent.")
        with open("text_output.txt", "w") as file:
            file.write("\n".join(context.user_data['text']))
        with open("text_output.txt", "rb") as file:
            update.message.reply_document(file)
        return ConversationHandler.END
    else:
        update.message.reply_text("Invalid format. Please choose one of txt, docx, or pdf.")
        return FORMAT

def get_font_size(update: Update, context: CallbackContext) -> int:
    try:
        font_size = int(update.message.text)
        if 10 <= font_size <= 20:
            context.user_data['font_size'] = font_size
            reply_keyboard = [['Arial', 'Times New Roman', 'Calibri', 'B Nazanin', 'B Mitra', 'B Yekan']]
            update.message.reply_text(
                "Font size received. Please choose a font style:",
                reply_markup=ReplyKeyboardMarkup(reply_keyboard, one_time_keyboard=True)
            )
            return FONT_STYLE
        else:
            update.message.reply_text("Please enter a number between 10 and 20.")
            return FONT_SIZE
    except ValueError:
        update.message.reply_text("Please enter a valid number for font size.")
        return FONT_SIZE

def get_font_style(update: Update, context: CallbackContext) -> int:
    context.user_data['font_style'] = update.message.text
    reply_keyboard = [['Bold', 'Italic', 'Bold & Italic', 'Normal']]
    update.message.reply_text(
        "Font style selected. Please choose a font weight (Bold, Italic, Bold & Italic, or Normal):",
        reply_markup=ReplyKeyboardMarkup(reply_keyboard, one_time_keyboard=True)
    )
    return FONT_BOLD_ITALIC

def get_font_bold_italic(update: Update, context: CallbackContext) -> int:
    style = update.message.text
    context.user_data['bold'] = 'Bold' in style
    context.user_data['italic'] = 'Italic' in style

    update.message.reply_text(
        "Do you want to edit the text? If yes, enter /edit. Otherwise, the file will be converted."
    )
    return EDIT

def edit_text(update: Update, context: CallbackContext) -> int:
    update.message.reply_text(
        "You are in edit mode.\n"
        "To delete a line, use the command /delete followed by the line number (e.g., /delete 3).\n"
        "To replace a word, use the command /replace followed by new_word old_word.\n"
        "To add a new line, use /add followed by the text.\n"
        "To finish and save changes, use /done."
    )
    return EDIT

def handle_edit_commands(update: Update, context: CallbackContext) -> int:
    command = update.message.text
    if command.startswith('/delete'):
        try:
            line_number = int(command.split()[1]) - 1
            if 0 <= line_number < len(context.user_data['text']):
                context.user_data['text'].pop(line_number)
                update.message.reply_text("Line deleted.")
            else:
                update.message.reply_text("Invalid line number.")
        except (IndexError, ValueError):
            update.message.reply_text("Please enter a valid line number.")
    
    elif command.startswith('/replace'):
        try:
            _, new_word, old_word = command.split()
            context.user_data['text'] = [line.replace(old_word, new_word) for line in context.user_data['text']]
            update.message.reply_text("Word replaced.")
        except ValueError:
            update.message.reply_text("Please enter the command as follows: /replace new_word old_word")

    elif command.startswith('/add'):
        new_line = command[5:]
        context.user_data['text'].append(new_line)
        update.message.reply_text("New line added.")

    elif command == '/done':
        if context.user_data['format'] == 'pdf':
            return convert_to_pdf(update, context)
        else:
            return convert_to_file(update, context)
    
    else:
        update.message.reply_text("Invalid command.")
    
    return EDIT

def convert_to_pdf(update: Update, context: CallbackContext) -> int:
    # ایجاد فایل PDF با استفاده از PyMuPDF
    doc = fitz.open()  # ایجاد یک سند PDF خالی
    text = "\n".join(context.user_data['text'])
    page = doc.new_page()  # اضافه کردن یک صفحه جدید
    rect = fitz.Rect(72, 72, 500, 800)  # تنظیم مکان متن در صفحه
    page.insert_textbox(rect, text, fontsize=12)  # درج متن در جعبه متن
    doc.save("text_output.pdf")  # ذخیره فایل PDF
    doc.close()

    # ارسال فایل PDF برای کاربر
    with open("text_output.pdf", "rb") as file:
        update.message.reply_document(file)
    
    update.message.reply_text("Your PDF file has been created successfully!")
    return ConversationHandler.END

def convert_to_file(update: Update, context: CallbackContext) -> int:
    # ایجاد فایل docx با تنظیمات قالب‌بندی
    if context.user_data['format'] == 'docx':
        doc = Document()
        for line in context.user_data['text']:
            para = doc.add_paragraph()
            run = para.add_run(line)
            run.font.size = Pt(context.user_data['font_size'])
            run.font.name = context.user_data['font_style']
            run.bold = context.user_data['bold']
            run.italic = context.user_data['italic']
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        doc.save("text_output.docx")
        with open("text_output.docx", "rb") as file:
            update.message.reply_document(file)
    else:
        with open("text_output.txt", "w") as file:
            file.write("\n".join(context.user_data['text']))
        with open("text_output.txt", "rb") as file:
            update.message.reply_document(file)

    update.message.reply_text("Your file has been created successfully!")
    return ConversationHandler.END

def cancel(update: Update, context: CallbackContext) -> int:
    update.message.reply_text(
        "Conversation canceled. If you want to try again, type /start."
    )
    return ConversationHandler.END

def main() -> None:
    updater = Updater(TOKEN)
    dispatcher = updater.dispatcher

    # تعریف ConversationHandler برای مراحل مکالمه
    conv_handler = ConversationHandler(
        entry_points=[CommandHandler("start", start)],
        states={
            TEXT: [MessageHandler(Filters.text & ~Filters.command, get_text)],
            FORMAT: [MessageHandler(Filters.regex('^(txt|docx|pdf)$'), get_format)],
            FONT_SIZE: [MessageHandler(Filters.text & ~Filters.command, get_font_size)],
            FONT_STYLE: [MessageHandler(Filters.text & ~Filters.command, get_font_style)],
            FONT_BOLD_ITALIC: [MessageHandler(Filters.text & ~Filters.command, get_font_bold_italic)],
            EDIT: [
                CommandHandler("edit", edit_text),
                MessageHandler(Filters.regex(r'^/delete \d+$'), handle_edit_commands),
                MessageHandler(Filters.regex(r'^/replace .+ .+$'), handle_edit_commands),
                MessageHandler(Filters.regex(r'^/add .+$'), handle_edit_commands),
                CommandHandler("done", convert_to_file)
            ],
        },
        fallbacks=[CommandHandler("cancel", cancel)],
    )

    dispatcher.add_handler(conv_handler)

    updater.start_polling()
    updater.idle()

if __name__ == '__main__':
    main()